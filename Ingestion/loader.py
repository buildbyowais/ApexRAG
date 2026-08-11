import os
from typing import List

from docx import Document as DocxDocument
from langchain_community.document_loaders import (
    PyMuPDFLoader,
    TextLoader
)
from langchain_core.documents import Document


def extract_docx(file_path: str) -> List[Document]:
    doc = DocxDocument(file_path)

    text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())

    content = "\n".join(text)

    return [
        Document(
            page_content=content,
            metadata={
                "source": file_path,
                "page": 1
            }
        )
    ]


def load_document(file_path: str) -> List[Document]:
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return PyMuPDFLoader(file_path).load()

    if extension == ".docx":
        return extract_docx(file_path)

    if extension == ".txt":
        documents = TextLoader(
            file_path,
            encoding="utf-8"
        ).load()

        for document in documents:
            document.metadata["page"] = 1

        return documents

    raise ValueError(
        "Unsupported file format. "
        "Supported formats: PDF, DOCX, TXT."
    )